import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ================================
# Page setup
# ================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
style.paragraph_format.line_spacing = 1.5


# ================================
# Helpers
# ================================
def cover_line(text, size=Pt(16), bold=True, gap=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = gap
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = size
    run.bold = bold


def blank(n=1):
    for _ in range(n):
        cover_line('', size=Pt(16), bold=False)


def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = 'SimSun'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'SimSun'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(12)


def table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(10)
    run.bold = True


def three_line_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(10)
        run.bold = True

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'SimSun'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(10)

    # Apply three-line borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '</w:tblBorders>' % nsdecls('w'))
    tblPr.append(borders)

    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders %s>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '</w:tcBorders>' % nsdecls('w'))
        tcPr.append(tcBorders)

    return table


# ================================
# COVER PAGE
# ================================
blank(3)
cover_line('华中科技大学', Pt(22))
cover_line('计算机科学与技术学院', Pt(18))
cover_line('《机器学习》结课报告', Pt(18))
blank(2)

for t in ['专    业：', '班    级：', '学    号：', '姓    名：', '成    绩：', '指导教师：']:
    cover_line(t, Pt(16), bold=False)

blank(1)
today = datetime.date.today()
cover_line('完成日期：%d 年 %d 月 %d 日' % (today.year, today.month, today.day), Pt(14), bold=False)
doc.add_page_break()

# ================================
# TOC
# ================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目  录')
run.font.size = Pt(16)
run.bold = True
run.font.name = 'SimSun'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

for e in ['1. 实验要求', '2. 算法设计与实现', '3. 实验环境与平台', '4. 结果与分析', '5. 个人体会']:
    p = doc.add_paragraph()
    run = p.add_run(e)
    run.font.size = Pt(14)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
doc.add_page_break()

# ================================
# 1. 实验要求
# ================================
h1('1. 实验要求')

h2('1.1 实验目的')
body('本实验旨在深入理解集成学习中 AdaBoost 算法的原理与实现。具体要求为：分别以对数几率回归（Logistic Regression）和决策树桩（Decision Stump）作为基分类器，从零实现 AdaBoost 算法，不允许调用任何已封装好的机器学习模型库。通过在不同基分类器数量（1、5、10、100）下进行 10 折交叉验证，对比分析基分类器类型和数量对集成模型性能的影响。')

h2('1.2 输入输出格式')
body('程序读取 data.csv（特征矩阵）和 targets.csv（标签）两个文件，输出预测结果至 experiments/baseX_foldY.csv。其中 X 为基分类器数量（1/5/10/100），Y 为测试折编号（1~10）。每个预测文件包含两列：第一列为样本序号（从 1 开始），第二列为预测标记（0 或 1）。程序入口 main.py 支持命令行调用：python main.py <data.csv> <targets.csv> <0|1>，其中 0 表示对数几率回归，1 表示决策树桩。')

# ================================
# 2. 算法设计与实现
# ================================
h1('2. 算法设计与实现')

h2('2.1 加权对数几率回归')
body('对数几率回归是一种线性分类模型，通过 sigmoid 函数 sigma(z) = 1/(1+e^{-z}) 将线性组合映射到 (0,1) 区间。为适配 AdaBoost 的样本权重机制，将标准交叉熵损失扩展为加权形式：L = -sum_i w_i [y_i log p_i + (1-y_i) log(1-p_i)]，其中 w_i 为归一化后的样本权重。在实现中，将偏置项并入权重矩阵（X 左侧拼接一列全 1），梯度下降的更新公式中除以样本数 m 以稳定训练。特征预先进行 z-score 标准化（均值与标准差在训练集上计算并应用于测试集，避免信息泄露）。主要超参数：学习率 0.01，最大迭代 1000 轮，不使用 L2 正则化。')

h2('2.2 决策树桩')
body('决策树桩是深度为 1 的决策树，仅使用单个特征进行一次划分。训练时对每个特征，按特征值 argsort 排序后遍历所有相邻样本对的中点作为候选阈值。对于每个阈值，同时评估两种极性（左侧 +1 右侧 -1，以及左侧 -1 右侧 +1），选择加权分类误差最小的（特征、阈值、极性）组合。该实现遍历了所有可能的信息增益分割点，确保找到全局最优的决策树桩。')

h2('2.3 AdaBoost 算法')
body('AdaBoost（Adaptive Boosting）通过迭代训练弱分类器并自适应调整样本权重，将多个弱分类器组合为强分类器。算法核心步骤如下：')
body('(1) 初始化所有样本权重为 1/N。')
body('(2) 每轮迭代在加权数据上训练基分类器 h_t，计算加权误差 epsilon_t = sum_i w_i * I[h_t(x_i) != y_i]。')
body('(3) 若 epsilon_t >= 0.5（差于随机猜测）或 epsilon_t = 0（完美分类），终止迭代。这一"止损"机制确保算法在无法继续提升时及时退出：若直接 continue 而不更新权重，因基分类器训练是确定性的（相同权重、相同数据得到相同结果），后续迭代将陷入无效的重复循环。')
body('(4) 计算分类器权重 alpha_t = 0.5 * ln((1-epsilon_t)/epsilon_t)。')
body('(5) 按 w_i <- w_i * exp(-alpha_t * y_i * h_t(x_i)) 更新样本权重并重新归一化。')
body('(6) 最终预测为各基分类器加权投票的符号：H(x) = sign(sum_t alpha_t * h_t(x))。')

h2('2.4 交叉验证设计')
body('采用 10 折交叉验证评估模型性能。使用固定随机种子（seed=42）将数据集随机打乱后均匀划分为 10 折（每折约 368 个样本），依次以其中 1 折为测试集、剩余 9 折（约 3312 个样本）为训练集。所有实验使用相同的折划分以保证结果可复现。特征标准化在每折的训练集上独立计算均值和标准差并应用于测试集，避免测试集信息泄露。')

# ================================
# 3. 实验环境与平台
# ================================
h1('3. 实验环境与平台')

table_caption('表 1  实验环境配置')
three_line_table(
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 11'],
        ['编程语言', 'Python 3.10'],
        ['依赖库', 'NumPy（仅用于矩阵运算）'],
        ['开发环境', 'VS Code'],
    ]
)
doc.add_paragraph()

table_caption('表 2  数据集信息')
three_line_table(
    ['属性', '值'],
    [
        ['样本总数', '3680'],
        ['特征维度', '57'],
        ['标签类型', '二分类（0/1）'],
        ['类别分布', '类别 0: 59.8%, 类别 1: 40.2%'],
        ['评估方式', '10 折交叉验证'],
    ]
)
doc.add_paragraph()
body('说明：本实验仅使用 NumPy 库进行矩阵运算，所有机器学习模型（对数几率回归、决策树桩、AdaBoost）均从零实现，未调用 scikit-learn 等任何机器学习库。')

# ================================
# 4. 结果与分析
# ================================
h1('4. 结果与分析')

h2('4.1 总体结果')
body('对两种基分类器在 4 种基分类器数量（T = 1, 5, 10, 100）下分别进行 10 折交叉验证，实验结果如表 3 所示。')

table_caption('表 3  不同基分类器数量下的 10 折交叉验证准确率')
three_line_table(
    ['基分类器数量 T', '对数几率回归', '决策树桩'],
    [
        ['1', '0.8984', '0.7913'],
        ['5', '0.9003', '0.8954'],
        ['10', '0.9052', '0.9062'],
        ['100', '0.9035', '0.9432'],
    ]
)
doc.add_paragraph()
body('两种基分类器在 T=100 时的准确率分别为 90.35% 和 94.32%，均超过 80% 的评测要求。下面分别对两组实验结果进行分析。')

h2('4.2 对数几率回归作为基分类器')
body('以对数几率回归为基分类器时，仅使用 1 个基分类器即可达到 89.84% 的准确率，说明该数据集具有较强的线性可分性。随着基分类器数量从 1 增加到 10，准确率从 89.84% 小幅上升至 90.52%（+0.68%），但继续增加到 100 时准确率回落至 90.35%。')
body('分析原因：对数几率回归本身是较强的线性分类器，第一个基分类器已经拟合了数据的主要线性结构。后续迭代中，被误分类的样本往往是线性不可分的"硬样本"，受限于线性决策边界，新的加权分类器难以有效纠正这些错误。AdaBoost 的核心优势在于组合大量弱分类器以降低偏差，而强分类器之间缺乏足够的差异性（diversity），导致集成增益有限，甚至可能因过拟合训练集的权重分布而轻微损害泛化性能。')

h2('4.3 决策树桩作为基分类器')
body('以决策树桩为基分类器的结果则完整展现了 AdaBoost 的核心优势。单个决策树桩仅有 79.13% 的准确率，属于典型的弱分类器，但随着基分类器数量的增加，性能呈现显著的持续提升趋势。')

table_caption('表 4  决策树桩 AdaBoost 各阶段性能提升')
three_line_table(
    ['阶段', '准确率变化', '提升幅度'],
    [
        ['T=1 -> T=5', '79.13% -> 89.54%', '+10.41%'],
        ['T=5 -> T=10', '89.54% -> 90.62%', '+1.08%'],
        ['T=10 -> T=100', '90.62% -> 94.32%', '+3.70%'],
        ['合计', '79.13% -> 94.32%', '+15.19%'],
    ]
)
doc.add_paragraph()
body('从表 4 可以看出，初期（T=1 到 5）提升最为显著，增益达 10.41%；中后期边际收益逐渐递减，但 T=100 时仍达到了 94.32% 的最高准确率，超越了以对数几率回归为基分类器的最佳结果（90.35%）。')
body('这一结果完美验证了 AdaBoost 的设计哲学：通过自适应调整样本权重，使后续分类器"关注"前序分类器出错的样本，将多个弱分类器（高偏差、低方差）逐级提升为强分类器。同时，决策树桩每次只选一个特征做一次划分，不同的桩可以选择不同的特征和分割点，因此 100 个桩的组合能够捕捉到单个线性模型无法表达的非线性决策边界，这正是 94.32% > 90.35% 的本质原因。')

h2('4.4 两种基分类器对比')

table_caption('表 5  两种基分类器的综合对比（T=100）')
three_line_table(
    ['对比维度', '对数几率回归', '决策树桩'],
    [
        ['单分类器强度', '强（89.84%）', '弱（79.13%）'],
        ['Boosting 提升空间', '小（+0.51%）', '大（+15.19%）'],
        ['最终准确率', '90.35%', '94.32%'],
        ['适合 AdaBoost', '不适合', '非常适合'],
        ['决策边界类型', '线性', '分段线性（可表达非线性）'],
    ]
)
doc.add_paragraph()
body('综上，AdaBoost 应搭配弱分类器使用才能发挥最大效用。决策树桩因其结构简单、训练高效、天然作为弱分类器，是 AdaBoost 最经典的基分类器选择。')

# ================================
# 5. 个人体会
# ================================
h1('5. 个人体会')

body('通过本次实验，我获得了以下几点深刻认识：')

body('第一，"弱可学习"是 AdaBoost 发挥效用的前提。实验数据直观地展示了这一结论：对数几率回归单模型已很强（89.84%），集成后几乎无提升；而决策树桩单模型很弱（79.13%），集成后反超强分类器达到 94.32%。AdaBoost 的核心机制是通过调整样本权重让后续分类器关注"难样本"，但这只有在基分类器有足够提升空间时才有效。强分类器做基分类器不仅浪费计算资源，还可能因多样性不足而损害泛化能力。')

body('第二，加权误差 epsilon_t >= 0.5 的"止损"机制十分关键。在实现初期使用了 continue 跳过该分类器，但忽略了权重未更新的问题。由于基分类器训练是确定性的（相同权重、相同数据得到相同结果），下一轮会陷入相同的跳过循环，造成大量无效计算。修改为 break 后，算法在无法继续提升时及时终止，既节省了计算资源，也更符合 AdaBoost 的理论设计。')

body('第三，Boosting 与 Bagging 的侧重点不同。决策树桩是高偏差、低方差的模型，AdaBoost 通过加权组合多个桩逐步降低整体偏差，拟合更复杂的决策边界。这与 Bagging（如随机森林）通过并行训练、平均化来降低方差的思想形成互补。理解了这一点，对集成学习家族的整体图景有了更清晰的把握。')

body('第四，超参数对模型性能有重要影响。实验中对比了不同参数设置下的对数几率回归表现：较大学习率（0.5）加少量迭代（200 轮）可获得约 91.9% 的准确率但无集成增益；较小学习率（0.01）加较多迭代（1000 轮）加梯度除以样本数则准确率约 90%，但表现出了一定的集成提升空间。这启示我们，在将强分类器用于 Boosting 时，适当削弱基分类器（减少迭代、增大正则化等）有可能增加多样性，带来集成增益。')

# ================================
# Save
# ================================
output_path = '机器学习结课报告_Adaboost算法实现.docx'
doc.save(output_path)
print('Done: ' + output_path)
